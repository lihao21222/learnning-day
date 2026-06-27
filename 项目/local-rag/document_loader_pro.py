import os
from typing import List, Dict, Any
from pathlib import Path
from dataclasses import dataclass
from abc import ABC, abstractmethod


@dataclass
class Document:
    """文档数据类"""
    content: str
    source: str
    file_type: str
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class BaseDocumentParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> List[Document]:
        """解析文档"""
        pass
    
    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """支持的文件扩展名"""
        pass


class TextParser(BaseDocumentParser):
    """文本文件解析器"""
    
    def parse(self, file_path: str) -> List[Document]:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return [Document(
                content=content,
                source=file_path,
                file_type='txt',
                metadata={'file_name': os.path.basename(file_path)}
            )]
        except Exception as e:
            print(f"解析 TXT 文件失败 {file_path}: {e}")
            return []
    
    def supported_extensions(self) -> List[str]:
        return ['.txt', '.md']


class PDFParser(BaseDocumentParser):
    """PDF 文件解析器"""
    
    def __init__(self):
        self.available = False
        self.pypdf = None
        self._init_parser()
    
    def _init_parser(self):
        try:
            import pypdf
            self.pypdf = pypdf
            self.available = True
        except ImportError:
            print("警告: pypdf 未安装，PDF 功能不可用")
            self.available = False
    
    def parse(self, file_path: str) -> List[Document]:
        if not self.available:
            print("PDF 解析器不可用，请安装: pip install pypdf")
            return []
        
        try:
            docs = []
            with open(file_path, 'rb') as f:
                reader = self.pypdf.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        docs.append(Document(
                            content=text,
                            source=file_path,
                            file_type='pdf',
                            metadata={
                                'file_name': os.path.basename(file_path),
                                'page': page_num + 1,
                                'total_pages': len(reader.pages)
                            }
                        ))
            return docs
        except Exception as e:
            print(f"解析 PDF 文件失败 {file_path}: {e}")
            return []
    
    def supported_extensions(self) -> List[str]:
        return ['.pdf']


class DocxParser(BaseDocumentParser):
    """Word 文档解析器"""
    
    def __init__(self):
        self.available = False
        self.docx = None
        self._init_parser()
    
    def _init_parser(self):
        try:
            from docx import Document as DocxDocument
            self.docx = DocxDocument
            self.available = True
        except ImportError:
            print("警告: python-docx 未安装，DOCX 功能不可用")
            self.available = False
    
    def parse(self, file_path: str) -> List[Document]:
        if not self.available:
            print("DOCX 解析器不可用，请安装: pip install python-docx")
            return []
        
        try:
            docs = []
            doc = self.docx(file_path)
            
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            if full_text:
                docs.append(Document(
                    content='\n'.join(full_text),
                    source=file_path,
                    file_type='docx',
                    metadata={
                        'file_name': os.path.basename(file_path)
                    }
                ))
            
            return docs
        except Exception as e:
            print(f"解析 DOCX 文件失败 {file_path}: {e}")
            return []
    
    def supported_extensions(self) -> List[str]:
        return ['.docx']


class DocumentLoaderPro:
    """增强版文档加载器"""
    
    def __init__(self, knowledge_base_dir: str = None):
        self.knowledge_base_dir = knowledge_base_dir or 'knowledge_base'
        self.parsers = self._init_parsers()
        self.extension_map = self._build_extension_map()
    
    def _init_parsers(self) -> List[BaseDocumentParser]:
        """初始化所有解析器"""
        return [
            TextParser(),
            PDFParser(),
            DocxParser()
        ]
    
    def _build_extension_map(self) -> Dict[str, BaseDocumentParser]:
        """建立扩展名到解析器的映射"""
        ext_map = {}
        for parser in self.parsers:
            for ext in parser.supported_extensions():
                ext_map[ext.lower()] = parser
        return ext_map
    
    def load_file(self, file_path: str) -> List[Document]:
        """加载单个文件"""
        _, ext = os.path.splitext(file_path.lower())
        
        if ext not in self.extension_map:
            print(f"不支持的文件格式: {ext}")
            return []
        
        parser = self.extension_map[ext]
        return parser.parse(file_path)
    
    def load_directory(self, directory: str = None) -> List[Document]:
        """加载整个目录"""
        dir_path = directory or self.knowledge_base_dir
        all_docs = []
        
        if not os.path.exists(dir_path):
            print(f"目录不存在: {dir_path}")
            return all_docs
        
        for filename in os.listdir(dir_path):
            file_path = os.path.join(dir_path, filename)
            
            if os.path.isfile(file_path):
                _, ext = os.path.splitext(filename.lower())
                
                if ext in self.extension_map:
                    print(f"正在加载: {filename}")
                    docs = self.load_file(file_path)
                    all_docs.extend(docs)
                else:
                    print(f"跳过不支持的文件: {filename}")
        
        print(f"目录加载完成: 总计 {len(all_docs)} 个文档")
        return all_docs
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的格式列表"""
        formats = set()
        for parser in self.parsers:
            formats.update(parser.supported_extensions())
        return sorted(list(formats))
